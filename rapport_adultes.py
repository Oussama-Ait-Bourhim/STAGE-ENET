import math
import numpy as np
import pylab as pl
import seaborn as sns
import pandas

#importation data
import pandas as pd
df=pd.read_excel("data\carnet_adulte_ENET2012.csv")


#data selon les noms techniques de nomoncaltures
df['Milieu']= df['Milieu'].replace([1,2],['urbain','rural'])
#df['Sexe']= df['Sexe'].replace([1,2],['hommes','femmes'])
df['Ramadan']= df['Ramadan'].replace([0,1],['hors ramadan','ramadan'])
df['Groupe_âge']= df['Groupe_âge'].replace([1,2,3,4,5],['15 à 24','25 à 34','35 à 45','46 à 59','60 ans au plus'])
df['Etat_matrimonial']= df['Etat_matrimonial'].replace([1,2,3,4],['celibataire','marié','divorcé','veuf'])
df['Niveau_scolaire']= df['Niveau_scolaire'].replace([0,1,2,3,4,5],['sans niveau','primaire','secondaire collegial','secondaire qualifiant','superieur','autre niveau'])
df['Type_activite']= df['Type_activite'].replace([1,2,3,4,5],['actif occupé','chomeur','femme foyer','etudiant','autres'])
df['Catégorie_socioprofessionnelle']= df['Catégorie_socioprofessionnelle'].replace([1,2,3,4,5,6,9],['directeur','cadre moyen','commerçants','exploitant','artisants','monoeuvres','non declaré'])
df['Statut_professionnel']= df['Statut_professionnel'].replace([1,2,3,9],['salarié','autoemployé','non remineré','non declaré'])




#definition des fonction necessaire
def meantime(df,activ,ind2,val2):
    df1=df[[activ, ind2]].loc[df[ind2]==val2]
    r=df1[activ].mean()
    return r


def meantime_sx(df,activ,ind2,val2,sexe):
    df1=df[[activ, ind2]].loc[(df[ind2]==val2) & (df['Sexe']==sexe)]
    r=df1[activ].mean()
    return r

def meantime_sxh(df,activ,ind2,val2):
    r=meantime_sx(df, activ, ind2, val2, 1)
    return r

def meantime_sxf(df,activ,ind2,val2):
    r=meantime_sx(df, activ, ind2, val2,2)
    return r

def list_time(n,ind2,val2):
    list_activ=['dureeG0','dureeG1','dureeG2','dureeG3','dureeG4','dureeG5','dureeG6','dureeG7','dureeG8','dureeG9']
    x=[]
    for i in range (0,len(list_activ)):
      r1=n(df,list_activ[i],ind2,val2)
      if pandas.isna(r1)==False:
          r=int(r1)
          x.append(r)
      else: x.append(r1)
    t_physo=x[0]+x[1]+x[2]
    tm=x[5]+x[6]
    t_libre=x[7]+x[8]+x[9]
    x.insert(0,t_physo)
    x.insert(6,tm)
    x.insert(9,t_libre)
    return x

def total_time():
    list=['dureeG0','dureeG1','dureeG2','dureeG3','dureeG4','dureeG5','dureeG6','dureeG7','dureeG8','dureeG9']
    z=[]
    for i in range (0,len(list)):
      r2=df[list[i]].mean()
      if pandas.isna(r2)==False:
        r=int(r2)
        z.append(r)
      else: z.append(r2)
    t_physo=z[0]+z[1]+z[2]
    tm=z[5]+z[6]
    t_libre=z[7]+z[8]+z[9]
    z.insert(0,t_physo)
    z.insert(6,tm)
    z.insert(9,t_libre)
    return z

def total_time_sx(sexe):
    list=['dureeG0','dureeG1','dureeG2','dureeG3','dureeG4','dureeG5','dureeG6','dureeG7','dureeG8','dureeG9']
    z=[]
    for i in range (0,len(list)):
      r2=df[list[i]].loc[(df['Sexe']==sexe)].mean()
      if pandas.isna(r2)==False:
        r=int(r2)
        z.append(r)
      else: z.append(r2)
    t_physo=z[0]+z[1]+z[2]
    tm=z[5]+z[6]
    t_libre=z[7]+z[8]+z[9]
    z.insert(0,t_physo)
    z.insert(6,tm)
    z.insert(9,t_libre)
    return z

def total_time_sxh():
    z=[]
    z=total_time_sx(1)
    return z

def total_time_sxf():
    z=[]
    z=total_time_sx(2)
    return z





def cont_data(listtime,mean,total,indic):
    col1 = ['Temps physiologique', 'Sommeil', 'Repas', 'Soins personnel', 'Travail professionnel',
            'Formation et éducation',
            'Travaux ménagers,soins du ménage', 'Travaux ménagers',
            'Soins donnés aux membres du ménage', 'Temps libre',
            'Loisirs',

            'Sociabilité', 'Pratiques religieuses']
    data = {'Activité': col1}
    d = pd.DataFrame(data)
    a = df[indic]
    c = list(set(a))
    for i in range(0, len(c)):
        x = []
        x = listtime(mean,indic, c[i])
        d.loc[:, c[i]] = x
    z = []
    z = total()
    d.loc[:, 'Ensemble'] = z
    return d

#tableau moyenne time par milieu
df_milieu=cont_data(list_time,meantime,total_time,'Milieu')
#tableau moyenne time par milieu pour homme
df_milieu_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Milieu')
#tableau moyenne time par milieu pour femme
df_milieu_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Milieu')


#tableau moyenne time par etat matrimoine
df_etatm=cont_data(list_time,meantime,total_time,'Etat_matrimonial')
#tableau moyenne time paretat pour homme
df_etatm_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Etat_matrimonial')
#tableau moyenne time par etat pour femme
df_etatm_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Etat_matrimonial')

#tableau moyenne time par groupe age
df_age=cont_data(list_time,meantime,total_time,'Groupe_âge')
#tableau moyenne time par  age pour homme
df_age_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Groupe_âge')
#tableau moyenne time par  age pour femme
df_age_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Groupe_âge')

#tableau moyenne time par niveau scolarité
df_etude=cont_data(list_time,meantime,total_time,'Niveau_scolaire')
#tableau moyenne time par niveau scolarité pour homme
df_etude_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Niveau_scolaire')
#tableau moyenne time par niveau scolarité pour femme
df_etude_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Niveau_scolaire')

#tableau moyenne time par rtipe activité
df_activite=cont_data(list_time,meantime,total_time,'Type_activite')
#tableau moyenne time paretat pour homme
df_activite_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Type_activite')
#tableau moyenne time par etat pour femme
df_activite_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Type_activite')

#tableau moyenne time par STATUT pROFEFESIONEL
df_profesion=cont_data(list_time,meantime,total_time,'Statut_professionnel')
#tableau moyenne time STATUT PROFEFESIONEL pour homme
df_profesion_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Statut_professionnel')
#tableau moyenne time STATUT PROFEFESIONEL pour femme
df_profesion_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Statut_professionnel')

#tableau moyenne time par STATUT socioPROFEFESIONEL
df_profesionsocio=cont_data(list_time,meantime,total_time,'Catégorie_socioprofessionnelle')
#tableau moyenne time STATUT socioPROFEFESIONEL pour homme
df_profesionsocio_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Catégorie_socioprofessionnelle')
#tableau moyenne time STATUT socioPROFEFESIONEL pour femme
df_profesionsocio_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Catégorie_socioprofessionnelle')

#tableau moyenne time par STATUT socioPROFEFESIONEL
df_ramadan=cont_data(list_time,meantime,total_time,'Ramadan')
#tableau moyenne time STATUT socioPROFEFESIONEL pour homme
df_ramadan_sxh=cont_data(list_time,meantime_sxh,total_time_sxh,'Ramadan')
#tableau moyenne time STATUT socioPROFEFESIONEL pour femme
df_ramdan_sxf=cont_data(list_time,meantime_sxf,total_time_sxf,'Ramadan')






import docx
import pandas as pd
doc = docx.Document()

def dataframe_to_word(data):
    df = pd.DataFrame(data)
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])

    t.style = 'Table Grid'
    doc.save('./rapport_adultes.docx')


doc.add_heading('Enquête nationale sur l’emploi du temps des marocains 2012', 0)
doc.add_heading('Tableaux statistiques pour les adultes (age +15 ans)', )
doc.add_page_break()

#selon Milieu
doc.add_heading('Tableaux1: Temps moyenne (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillées, le milieu de résidence et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_milieu)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_milieu_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_milieu_sxf)
doc.add_page_break()
#Etat Matrimonial
doc.add_heading('Tableau 2 : Temps moyen (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillées le groupe âge et le sexe ', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_age)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_age_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_age_sxf)
doc.add_page_break()

#Etat Matrimonial
doc.add_heading('Tableau 3 : Temps moyen (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillés, l’état matrimonial et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_etatm)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_etatm_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_etatm_sxf)
doc.add_page_break()

#selon niveau scolaire
doc.add_heading('Tableau 4 : Temps moyen (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillés, le  niveau scolaire et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_etude)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_etude_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_etude_sxf)
doc.add_page_break()
#selon type activité
doc.add_heading('Tableau 5 : Temps moyen (en minute par jour) de la population de âgée 15 ans et plus selon les activités détaillés, le type activité et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_activite)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_activite_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_activite_sxf)
doc.add_page_break()
#selonproffesion
doc.add_heading('Tableaux 6 :Temps moyen (en minute par jour) de la population active occupée âgée de 15 ans et plus selon les activités détaillés, la situation dans emploi et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_profesion)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_profesion_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_profesion_sxf)
doc.add_page_break()
#selon profession_socio
doc.add_heading('Tableau 7 : Temps moyen (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillées, la catégorie socio professionnelle et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_profesionsocio)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_profesionsocio_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_profesionsocio_sxf)
doc.add_page_break()
#selon ramadan
doc.add_heading('Tableau 8 : Temps moyen (en minute par jour) de la population âgée de 15 ans et plus selon les activités détaillés, le type de mois et le sexe', 0)
doc.add_heading('ENSEMBLE', 3)
dataframe_to_word(df_ramadan)
doc.add_heading('HOMMES', 3)
dataframe_to_word(df_ramadan_sxh)
doc.add_heading('FEMMMES', 3)
dataframe_to_word(df_ramdan_sxf)








